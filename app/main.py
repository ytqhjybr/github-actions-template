import os
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
from docx import Document
from app.services.llm_service import get_llm_response
from app.db.init_db import init_db
from app.services.proposal_service import save_proposal
from app.services.proposal_generator import generate_proposal_text
from app.services.rag_service import index_pdf, search_pdf
from app.services.purchase_assistant import analyze_stock
from app.services.analytics import get_orders_data, get_proposals_data

load_dotenv()

app = FastAPI(title="AI Assistants for Agricultural Dealer")

# Разрешаем запросы из Streamlit (порт 8501)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501", "http://127.0.0.1:8501"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()

# ========== ОБЩИЕ НАСТРОЙКИ ==========
UPLOAD_DIR = "data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ========== ЭНДПОИНТЫ ==========

@app.get("/")
async def root():
    return {"message": "API is running"}

@app.get("/health")
async def health():
    return {"status": "ok"}

@app.post("/test-llm")
async def test_llm(prompt: str):
    if get_llm_response:
        response = get_llm_response(prompt)
        return {"prompt": prompt, "response": response}
    return {"prompt": prompt, "response": "LLM сервис не доступен"}

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    return {"filename": file.filename, "path": file_path, "status": "uploaded"}

class ProposalRequest(BaseModel):
    client_inn: str
    region: str
    specialization: str
    additional_params: dict = None

@app.post("/generate-proposal")
async def generate_proposal(request: ProposalRequest):
    file_path = f"data/proposals/proposal_{request.client_inn}.docx"
    proposal_id = save_proposal(request.client_inn, file_path) if save_proposal else None
    return {
        "status": "generated",
        "proposal_id": proposal_id,
        "client_inn": request.client_inn,
        "region": request.region,
        "specialization": request.specialization,
        "additional_params": request.additional_params,
        "file_path": file_path,
        "message": "Заглушка: КП сгенерировано и сохранено в БД. Реальный файл пока не создан."
    }

@app.post("/create-test-doc")
async def create_test_doc():
    doc = Document()
    doc.add_heading('Тестовый документ', level=1)
    doc.add_paragraph('Этот файл создан с помощью python-docx.')
    doc.add_paragraph('Можно использовать для проверки работы библиотеки.')

    os.makedirs("data/proposals", exist_ok=True)
    filename = "test_document.docx"
    file_path = os.path.join("data/proposals", filename)
    doc.save(file_path)

    return {
        "status": "created",
        "file": file_path,
        "message": f"Файл {filename} успешно создан в папке data/proposals"
    }

@app.post("/generate-from-template")
async def generate_from_template(
    template_name: str,
    client_name: str = "Клиент",
    region: str = "Регион",
    specialization: str = "Специализация",
    price_list_name: str = None,
    additional_params: dict = None
):
    template_path = os.path.join(UPLOAD_DIR, template_name)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    price_list_path = None
    if price_list_name:
        price_list_path = os.path.join(UPLOAD_DIR, price_list_name)
        if not os.path.exists(price_list_path):
            raise HTTPException(status_code=404, detail="Прайс-лист не найден")

    proposal_text = "Текст КП (заглушка)"
    if generate_proposal_text:
        proposal_text = generate_proposal_text(
            region=region,
            specialization=specialization,
            price_list_path=price_list_path,
            additional_params=additional_params
        )

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        if "{{client_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{client_name}}", client_name)
        if "{{region}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{region}}", region)
        if "{{specialization}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{specialization}}", specialization)
        if "{{proposal_text}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{proposal_text}}", proposal_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{client_name}}" in cell.text:
                    cell.text = cell.text.replace("{{client_name}}", client_name)
                if "{{region}}" in cell.text:
                    cell.text = cell.text.replace("{{region}}", region)
                if "{{specialization}}" in cell.text:
                    cell.text = cell.text.replace("{{specialization}}", specialization)
                if "{{proposal_text}}" in cell.text:
                    cell.text = cell.text.replace("{{proposal_text}}", proposal_text)

    os.makedirs("data/proposals", exist_ok=True)
    output_filename = f"proposal_{client_name}.docx"
    output_path = os.path.join("data/proposals", output_filename)
    doc.save(output_path)

    return {
        "status": "generated",
        "output_file": output_path,
        "message": f"КП создано из шаблона {template_name}"
    }

class OrderRequest(BaseModel):
    vin: str
    model: str
    year: str
    description: str
    photo_path: Optional[str] = None
    region: str = "Московская область"
    specialization: str = "растениеводство"

@app.post("/save-order")
async def save_order(order: OrderRequest):
    import json
    from datetime import datetime
    os.makedirs("data/orders", exist_ok=True)
    filename = f"order_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    file_path = os.path.join("data/orders", filename)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(order.model_dump_json(indent=2, ensure_ascii=False))
    return {"status": "ok", "filename": filename}

from fastapi.responses import FileResponse

@app.get("/download/{file_path:path}")
async def download_file(file_path: str):
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Файл не найден")
    return FileResponse(file_path, filename=os.path.basename(file_path))

@app.post("/upload-pdf/")
async def upload_pdf(file: UploadFile = File(...)):
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Только PDF файлы")
    
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    if index_pdf:
        try:
            chunks_count = index_pdf(file_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ошибка индексации: {str(e)}")
    else:
        chunks_count = 0
    
    return {
        "status": "indexed",
        "filename": file.filename,
        "chunks": chunks_count
    }

@app.post("/ask-rag")
async def ask_rag(query: str):
    if not search_pdf:
        return {
            "query": query,
            "answer": "RAG сервис не доступен. Проверьте установку langchain и sentence-transformers."
        }
    
    fragments = search_pdf(query)
    
    if not fragments:
        return {
            "query": query,
            "answer": "Ничего не найдено в загруженных документах. Пожалуйста, загрузите руководства по ремонту или эксплуатации."
        }
    
    context = "\n\n".join(fragments)
    prompt = f"""Ответь на вопрос, используя только информацию из предоставленных фрагментов документов.
Если ответа нет в документах, скажи, что информации недостаточно.

Вопрос: {query}

Фрагменты документов:
{context}

Ответ:"""
    
    if get_llm_response:
        try:
            answer = get_llm_response(prompt, system_prompt="Ты ассистент сервисного инженера. Отвечай только по предоставленным документам.")
        except Exception as e:
            answer = f"Ошибка LLM: {e}\n\nНайденные фрагменты:\n{context}"
    else:
        answer = f"LLM не доступен. Найденные фрагменты:\n{context}"
    
    return {
        "query": query,
        "found_chunks": fragments,
        "answer": answer
    }

# ========== МОДУЛЬ 4: АССИСТЕНТ ЗАКУПЩИКА ==========
@app.post("/analyze-stock")
async def analyze_stock_endpoint(
    stock_file: UploadFile = File(...),
    orders_file: Optional[UploadFile] = None,
    in_transit_file: Optional[UploadFile] = None,
    sales_file: Optional[UploadFile] = None,
    safety_stock_formula: str = "mean_sales * 2 + min_stock"
):
    stock_path = os.path.join(UPLOAD_DIR, stock_file.filename)
    with open(stock_path, "wb") as f:
        f.write(await stock_file.read())
    
    orders_path = None
    if orders_file:
        orders_path = os.path.join(UPLOAD_DIR, orders_file.filename)
        with open(orders_path, "wb") as f:
            f.write(await orders_file.read())
    
    in_transit_path = None
    if in_transit_file:
        in_transit_path = os.path.join(UPLOAD_DIR, in_transit_file.filename)
        with open(in_transit_path, "wb") as f:
            f.write(await in_transit_file.read())
    
    sales_path = None
    if sales_file:
        sales_path = os.path.join(UPLOAD_DIR, sales_file.filename)
        with open(sales_path, "wb") as f:
            f.write(await sales_file.read())
    
    result = analyze_stock(
        stock_path,
        orders_path,
        in_transit_path,
        sales_path,
        safety_stock_formula
    )
    
    return result

# ========== ДАШБОРД ==========
@app.get("/analytics")
async def get_analytics():
    orders = get_orders_data()
    proposals = get_proposals_data()
    
    stats = {
        "total_orders": len(orders),
        "total_proposals": len(proposals),
        "orders_by_region": orders.groupby("region").size().to_dict() if not orders.empty else {},
        "orders_by_specialization": orders.groupby("specialization").size().to_dict() if not orders.empty else {},
        "proposals_by_client": proposals.groupby("client_name").size().to_dict() if not proposals.empty else {}
    }
    return stats
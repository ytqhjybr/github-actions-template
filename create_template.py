from docx import Document

doc = Document()
doc.add_paragraph("Коммерческое предложение для {{client_name}}")
doc.add_paragraph("Регион: {{region}}")
doc.add_paragraph("Специализация: {{specialization}}")
doc.add_paragraph("С уважением, ваша компания.")

doc.save("data/uploads/template.docx")
print("Шаблон успешно создан")
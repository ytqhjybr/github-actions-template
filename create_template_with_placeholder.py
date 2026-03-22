from docx import Document

doc = Document()
doc.add_paragraph("Коммерческое предложение")
doc.add_paragraph("Клиент: {{client_name}}")
doc.add_paragraph("Регион: {{region}}")
doc.add_paragraph("Специализация: {{specialization}}")
doc.add_paragraph("\n=== Текст предложения ===\n{{proposal_text}}")
doc.add_paragraph("\nС уважением, ваша компания.")

doc.save("data/uploads/template.docx")
print("Шаблон создан с плейсхолдером {{proposal_text}}")